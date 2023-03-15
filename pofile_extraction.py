from docx import *
import re
from datetime import datetime
class pofile_extraction():
    def __init__(self,filename):
        self.pofile_data={}
        self.document = Document(filename)
        self.get_date()
        self.get_po_number()
        self.get_file_order_number()
        self.get_address()
        self.get_contact()
        self.get_subject()
        self.get_total()
        
    def get_date(self):
        dates=[]
        date_pattern = "\d{2}[/-]\d{2}[/-]\d{4}"
        for x in  self.document.paragraphs:
            dates=re.findall(date_pattern,x.text)
            if(len(dates)!=0):
                break
        l=list(map(int,dates[0].split("/")))
        dat=datetime(l[2],l[1],l[0])
        self.pofile_data["Date"]=dat
    
    def get_po_number(self):
        po_pattern="STPIN/PUR/PO/{1}\d{2}[-]\d{2}[/]\d{2}"
        po=[]
        for x in  self.document.paragraphs:
            po=re.findall(po_pattern,x.text)
            
            if(len(po)!=0):
                break
        self.pofile_data["Purchase_Order_Number"]=po[0]
    
    def get_file_order_number(self):
        file_order_pattern="[A-Z]{4}[/][A-Z]{3}[^\s]*"
        file_order=[]
        for x in  self.document.paragraphs:
            file_order=re.findall(file_order_pattern,x.text)
            if(len(file_order)!=0):
                break
        self.pofile_data["File_Order_Number"]=file_order[0]

    def get_address(self):
        address_pattern="M/s"
        address_pattern2="Contact"
        address=[]
        for x in self.document.paragraphs:
    
            if( re.search(address_pattern,x.text)):
                # print(address)
                address=""
            if(re.search(address_pattern2,x.text)):
                break
            address+=x.text
            # print(address)
        self.pofile_data["Address"]=address
    def get_contact(self):
        contact_pattern="\d{10}"
        contact=[]
        for x in  self.document.paragraphs:
            contact=re.findall(contact_pattern,x.text)
            
            if(len(contact)!=0):
                break
        self.pofile_data["Contact"]=contact[0]
        
    def get_subject(self):
        subject_pattern="Subject.*"
        subject=[]
        for x in  self.document.paragraphs:
            subject=re.findall(subject_pattern,x.text)
            
            if(len(subject)!=0):
                break
        subject=subject[0].replace("Subject:","")
        self.pofile_data["Subject"]=subject
    
    def get_total(self):
        total_pattern="(?i)grand total"
        total=0 
        data=[]
        for table in self.document.tables:
            for i, row in enumerate(table.rows):
                text = (cell.text for cell in row.cells)


                row_data = tuple(text)
                data.append(row_data)
        
        for x in data:
            for y in x:
                
                if(re.search(total_pattern,y)):
                    total=x[-1]
        total=total.replace(',','')
        self.pofile_data["Amount"]=float(total)
    def return_data(self):
        return self.pofile_data
    def display(self):
        print(self.pofile_data)
# obj=pofile_extraction("pofile.docx")
# obj.display()