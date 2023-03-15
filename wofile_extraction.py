from docx import *
import re
from datetime import datetime
class wofile_extraction():
    def __init__(self,filename):
        self.wofile_data={}
        self.document = Document(filename)
        self.get_date()
        self.get_po_number()
        self.get_file_order_number()
        self.get_address()
        self.get_contact()
        self.get_subject()
        self.get_total()
        self.get_period()
    def get_date(self):
        dates=[]
        date_pattern = "\d{2}[/-]\d{2}[/-]\d{4}"
        for x in  self.document.paragraphs:
            dates=re.findall(date_pattern,x.text)
            if(len(dates)!=0):
                break
        l=list(map(int,dates[0].split("/")))
        dat=datetime(l[2],l[1],l[0])
        self.wofile_data["Date"]=dat
    
    def get_po_number(self):
        wo_pattern="STPIN/PUR/WO/{1}\d{2}[-]\d{2}[/]\d{2}"
        wo=[]
        for x in  self.document.paragraphs:
            po=re.findall(wo_pattern,x.text)
            
            if(len(po)!=0):
                break
        self.wofile_data["Work_Order_Number"]=po[0]
    
    def get_file_order_number(self):
        file_order_pattern="[A-Z]{4}[/][A-Z]{3}[^\s]*"
        file_order=[]
        for x in  self.document.paragraphs:
            file_order=re.findall(file_order_pattern,x.text)
            if(len(file_order)!=0):
                break
        self.wofile_data["File_Order_Number"]=file_order[0]

    def get_address(self):
        address_pattern="M/s"
        address_pattern2="Contact"
        address=[]
        for x in self.document.paragraphs:
    
            if( re.search(address_pattern,x.text)):
                # print(address)
                address=""
            if(re.search(address_pattern2,x.text)):
                break
            address+=x.text
            # print(address)
        self.wofile_data["Address"]=address
    def get_contact(self):
        contact_pattern="\d{10}"
        contact=[]
        for x in  self.document.paragraphs:
            contact=re.findall(contact_pattern,x.text)
            
            if(len(contact)!=0):
                break
        self.wofile_data["Contact"]=contact[0]
        
    def get_subject(self):
        subject_pattern="Subject.*"
        subject=[]
        for x in  self.document.paragraphs:
            subject=re.findall(subject_pattern,x.text)
            
            if(len(subject)!=0):
                break
        subject=subject[0].replace("Subject:","")
        self.wofile_data["Subject"]=subject
    
    def get_total(self):
        total_pattern="(?i)grand total"
        total=0 
        data=[]
        for table in self.document.tables:
            for i, row in enumerate(table.rows):
                text = (cell.text for cell in row.cells)


                row_data = tuple(text)
                data.append(row_data)
        
        for x in data:
            for y in x:
                
                if(re.search(total_pattern,y)):
                    total=x[-1]
        total=total.replace(',','')
        self.wofile_data["Amount"]=float(total)
    def get_period(self):
        data = []
        start_period_pattern="(?i)AMC Start Period"
        end_period_pattern="(?i)AMC End Period"
        period=[]
        # keys = None
        for table in self.document.tables:
            for i, row in enumerate(table.rows):
                text = (cell.text for cell in row.cells)
                row_data = tuple(text)
                data.append(row_data)
        
        for x in data:
            for y in x:
                
                if(re.search(start_period_pattern,y)):
                    period.append(x[-1])
                elif(re.search(end_period_pattern, y)):
                    period.append(x[-1])
        l=list(map(int,period[0].split("/")))
        
        self.wofile_data["AMC_Start_Period"]=datetime(l[2],l[1],l[0]) 
        l=list(map(int,period[1].split("/")))
        self.wofile_data["AMC_End_Period"]=datetime(l[2],l[1],l[0]) 
        
    def return_data(self):
        return self.wofile_data
        
        
    def display(self):
        print(self.wofile_data)
# obj=wofile_extraction("wofile.docx")
# obj.display()